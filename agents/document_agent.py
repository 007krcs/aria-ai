"""
ARIA — Document Agent
======================
Full-spectrum document reading, writing, and analysis agent.

Supported formats:
    .docx  .xlsx  .csv  .pdf  .txt  .md  .json  .py  .js

Every public method returns:
    {"ok": bool, "result": str, "data": any}

Usage:
    from agents.document_agent import DocumentAgent
    agent = DocumentAgent()
    out = agent.read_docx("report.docx")
"""

import os
import re
import csv
import json
import mimetypes
from pathlib import Path
from datetime import datetime
from typing import Any, Optional, Union


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _ok(result: str = "", data: Any = None) -> dict:
    return {"ok": True, "result": result, "data": data}


def _err(result: str = "", data: Any = None) -> dict:
    return {"ok": False, "result": result, "data": data}


def _ext(path: str) -> str:
    return Path(path).suffix.lower()


SUPPORTED = {".docx", ".xlsx", ".csv", ".pdf", ".txt", ".md",
             ".json", ".py", ".js"}

TEXT_LIKE = {".txt", ".md", ".json", ".py", ".js"}


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT AGENT
# ─────────────────────────────────────────────────────────────────────────────

class DocumentAgent:
    """Reads, writes, analyses and converts documents for ARIA."""

    # ── Word (.docx) ──────────────────────────────────────────────────────────

    def read_docx(self, path: str) -> dict:
        """Return full text and heading structure from a Word document."""
        try:
            from docx import Document  # python-docx
        except ImportError:
            return _err("python-docx not installed. Run: pip install python-docx")

        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")
        try:
            doc = Document(path)
            paragraphs = []
            headings = []
            full_text_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                full_text_parts.append(text)
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    level = re.search(r"\d+", style_name)
                    level = int(level.group()) if level else 1
                    headings.append({"level": level, "text": text})
                paragraphs.append(text)

            full_text = "\n".join(full_text_parts)
            data = {
                "full_text": full_text,
                "headings": headings,
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "word_count": len(full_text.split()),
            }
            return _ok(f"Read {len(paragraphs)} paragraphs, {len(headings)} headings.", data)
        except Exception as e:
            return _err(f"Error reading .docx: {e}")

    def write_docx(self, path: str, content: str, title: str = "") -> dict:
        """Create or overwrite a Word document with proper formatting."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return _err("python-docx not installed. Run: pip install python-docx")

        path = str(Path(path).expanduser())
        try:
            doc = Document()
            if title:
                heading = doc.add_heading(title, level=0)
                heading.style.font.size = Pt(18)

            for block in content.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                # Detect markdown-style headings
                if block.startswith("# "):
                    doc.add_heading(block[2:].strip(), level=1)
                elif block.startswith("## "):
                    doc.add_heading(block[3:].strip(), level=2)
                elif block.startswith("### "):
                    doc.add_heading(block[4:].strip(), level=3)
                else:
                    doc.add_paragraph(block)

            doc.save(path)
            return _ok(f"Document saved to {path}", {"path": path})
        except Exception as e:
            return _err(f"Error writing .docx: {e}")

    def append_docx(self, path: str, text: str) -> dict:
        """Append text to an existing Word document."""
        try:
            from docx import Document
        except ImportError:
            return _err("python-docx not installed. Run: pip install python-docx")

        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")
        try:
            doc = Document(path)
            doc.add_paragraph(text)
            doc.save(path)
            return _ok(f"Appended text to {path}", {"path": path})
        except Exception as e:
            return _err(f"Error appending to .docx: {e}")

    # ── Excel (.xlsx) ─────────────────────────────────────────────────────────

    def read_xlsx(self, path: str, sheet: Optional[str] = None) -> dict:
        """Read Excel file; return rows as list of dicts keyed by header."""
        try:
            import openpyxl
        except ImportError:
            return _err("openpyxl not installed. Run: pip install openpyxl")

        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            target = sheet if sheet and sheet in wb.sheetnames else wb.sheetnames[0]
            ws = wb[target]

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return _ok("Sheet is empty.", {"rows": [], "sheet": target})

            headers = [str(h) if h is not None else f"col_{i}"
                       for i, h in enumerate(rows[0])]
            data_rows = []
            for row in rows[1:]:
                record = {headers[i]: row[i] for i in range(len(headers))}
                data_rows.append(record)

            wb.close()
            return _ok(
                f"Read {len(data_rows)} rows from sheet '{target}'.",
                {"rows": data_rows, "sheet": target, "headers": headers}
            )
        except Exception as e:
            return _err(f"Error reading .xlsx: {e}")

    def write_xlsx(self, path: str, data: list, sheet_name: str = "Sheet1") -> dict:
        """Write list-of-dicts to Excel. Writes headers from first dict's keys."""
        try:
            import openpyxl
        except ImportError:
            return _err("openpyxl not installed. Run: pip install openpyxl")

        path = str(Path(path).expanduser())
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name

            if not data:
                wb.save(path)
                return _ok(f"Wrote empty workbook to {path}", {"path": path})

            headers = list(data[0].keys())
            ws.append(headers)
            for row in data:
                ws.append([row.get(h) for h in headers])

            wb.save(path)
            return _ok(f"Wrote {len(data)} rows to {path}", {"path": path, "rows": len(data)})
        except Exception as e:
            return _err(f"Error writing .xlsx: {e}")

    # ── CSV ───────────────────────────────────────────────────────────────────

    def read_csv(self, path: str) -> dict:
        """Read CSV; return list of dicts (DictReader) plus raw rows."""
        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")
        try:
            rows = []
            with open(path, newline="", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames or []
                for row in reader:
                    rows.append(dict(row))
            return _ok(
                f"Read {len(rows)} rows from CSV.",
                {"rows": rows, "headers": list(headers)}
            )
        except Exception as e:
            return _err(f"Error reading CSV: {e}")

    def write_csv(self, path: str, data: list) -> dict:
        """Write list-of-dicts to CSV."""
        path = str(Path(path).expanduser())
        try:
            if not data:
                with open(path, "w", newline="", encoding="utf-8") as f:
                    pass
                return _ok(f"Wrote empty CSV to {path}", {"path": path})

            headers = list(data[0].keys())
            with open(path, "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                writer.writerows(data)
            return _ok(f"Wrote {len(data)} rows to {path}", {"path": path})
        except Exception as e:
            return _err(f"Error writing CSV: {e}")

    # ── PDF ───────────────────────────────────────────────────────────────────

    def read_pdf(self, path: str) -> dict:
        """Extract text from PDF using pdfplumber (preferred) or pypdf."""
        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")

        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    pages.append({"page": i + 1, "text": text})
                full_text = "\n\n".join(p["text"] for p in pages if p["text"])
            return _ok(
                f"Extracted text from {len(pages)} pages.",
                {"full_text": full_text, "pages": pages, "page_count": len(pages)}
            )
        except ImportError:
            pass
        except Exception as e:
            return _err(f"pdfplumber error: {e}")

        # Fallback: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({"page": i + 1, "text": text})
            full_text = "\n\n".join(p["text"] for p in pages if p["text"])
            return _ok(
                f"Extracted text from {len(pages)} pages (pypdf).",
                {"full_text": full_text, "pages": pages, "page_count": len(pages)}
            )
        except ImportError:
            pass
        except Exception as e:
            return _err(f"pypdf error: {e}")

        return _err("No PDF library found. Run: pip install pdfplumber  OR  pip install pypdf")

    # ── Plain text / text-like ────────────────────────────────────────────────

    def read_txt(self, path: str) -> dict:
        """Read any text file (.txt, .md, .json, .py, .js, etc.)."""
        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")
        try:
            for encoding in ("utf-8", "utf-8-sig", "latin-1"):
                try:
                    with open(path, "r", encoding=encoding) as f:
                        content = f.read()
                    lines = content.splitlines()
                    return _ok(
                        f"Read {len(lines)} lines ({len(content)} chars).",
                        {"content": content, "lines": lines, "line_count": len(lines)}
                    )
                except UnicodeDecodeError:
                    continue
            return _err("Could not decode file with utf-8 or latin-1.")
        except Exception as e:
            return _err(f"Error reading file: {e}")

    def write_txt(self, path: str, content: str) -> dict:
        """Write plain text to any file path."""
        path = str(Path(path).expanduser())
        try:
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            return _ok(f"Wrote {len(content)} chars to {path}", {"path": path})
        except Exception as e:
            return _err(f"Error writing file: {e}")

    # ── Analysis ──────────────────────────────────────────────────────────────

    def summarize_document(self, path: str) -> dict:
        """
        Read any supported document and produce a 3–5 sentence summary.
        Uses basic extractive summarisation (no model required).
        """
        raw = self._read_any(path)
        if not raw["ok"]:
            return raw

        text = self._extract_text(raw)
        if not text:
            return _err("No text content found in document.")

        # Sentence-level extractive summary
        sentences = re.split(r"(?<=[.!?])\s+", text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

        if not sentences:
            return _ok("Document appears to contain no full sentences.", {"summary": ""})

        # Pick ~5 evenly spaced sentences
        if len(sentences) <= 5:
            chosen = sentences
        else:
            step = len(sentences) / 5
            chosen = [sentences[int(i * step)] for i in range(5)]

        summary = " ".join(chosen)
        return _ok("Summary generated.", {"summary": summary, "sentence_count": len(chosen)})

    def search_in_document(self, path: str, query: str) -> dict:
        """Find query text (case-insensitive) in document; return matching lines."""
        raw = self._read_any(path)
        if not raw["ok"]:
            return raw

        text = self._extract_text(raw)
        if not text:
            return _err("No text to search.")

        lines = text.splitlines()
        q_lower = query.lower()
        matches = []
        for i, line in enumerate(lines, start=1):
            if q_lower in line.lower():
                matches.append({"line": i, "text": line.strip()})

        if matches:
            return _ok(f"Found {len(matches)} match(es) for '{query}'.", {"matches": matches})
        return _ok(f"No matches found for '{query}'.", {"matches": []})

    def get_document_info(self, path: str) -> dict:
        """Return metadata: size, pages/rows/words, last modified, type."""
        path = str(Path(path).expanduser())
        if not os.path.exists(path):
            return _err(f"File not found: {path}")

        stat = os.stat(path)
        ext = _ext(path)
        size_bytes = stat.st_size
        modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")

        info: dict = {
            "path": path,
            "extension": ext,
            "size_bytes": size_bytes,
            "size_kb": round(size_bytes / 1024, 2),
            "last_modified": modified,
        }

        raw = self._read_any(path)
        if raw["ok"]:
            text = self._extract_text(raw)
            info["word_count"] = len(text.split()) if text else 0
            info["char_count"] = len(text) if text else 0

            d = raw.get("data") or {}
            if "page_count" in d:
                info["pages"] = d["page_count"]
            if "rows" in d:
                info["rows"] = len(d["rows"])
            if "headings" in d:
                info["headings"] = len(d["headings"])

        return _ok(f"Info for {Path(path).name}", info)

    def convert_to_txt(self, path: str) -> dict:
        """Convert any supported format to plain text; return the text."""
        raw = self._read_any(path)
        if not raw["ok"]:
            return raw
        text = self._extract_text(raw)
        return _ok(f"Converted to plain text ({len(text)} chars).", {"text": text})

    def list_recent_documents(self, max_files: int = 20) -> dict:
        """
        List recently modified documents from Desktop, Documents, Downloads.
        Returns files sorted by modification time (newest first).
        """
        home = Path.home()
        search_dirs = [
            home / "Desktop",
            home / "Documents",
            home / "Downloads",
        ]

        found = []
        for d in search_dirs:
            if not d.exists():
                continue
            for f in d.iterdir():
                if f.is_file() and f.suffix.lower() in SUPPORTED:
                    try:
                        mtime = f.stat().st_mtime
                        found.append({
                            "path": str(f),
                            "name": f.name,
                            "extension": f.suffix.lower(),
                            "modified": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S"),
                            "size_kb": round(f.stat().st_size / 1024, 2),
                            "_mtime": mtime,
                        })
                    except Exception:
                        continue

        found.sort(key=lambda x: x["_mtime"], reverse=True)
        for item in found:
            del item["_mtime"]

        recent = found[:max_files]
        return _ok(f"Found {len(recent)} recent documents.", {"documents": recent})

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _read_any(self, path: str) -> dict:
        """Dispatch to the correct reader based on file extension."""
        path = str(Path(path).expanduser())
        ext = _ext(path)

        if ext == ".docx":
            return self.read_docx(path)
        elif ext == ".xlsx":
            return self.read_xlsx(path)
        elif ext == ".csv":
            return self.read_csv(path)
        elif ext == ".pdf":
            return self.read_pdf(path)
        elif ext in TEXT_LIKE or ext == ".txt":
            return self.read_txt(path)
        else:
            return _err(f"Unsupported format: {ext}. Supported: {', '.join(sorted(SUPPORTED))}")

    def _extract_text(self, raw: dict) -> str:
        """Pull a unified text string from any _read_any result."""
        d = raw.get("data") or {}
        if isinstance(d, dict):
            for key in ("full_text", "content", "text"):
                if key in d and isinstance(d[key], str):
                    return d[key]
            # CSV/xlsx: flatten rows
            if "rows" in d:
                rows = d["rows"]
                if not rows:
                    return ""
                if isinstance(rows[0], dict):
                    lines = [" | ".join(str(v) for v in row.values()) for row in rows]
                else:
                    lines = [" | ".join(str(v) for v in row) for row in rows]
                headers = d.get("headers", [])
                header_line = " | ".join(str(h) for h in headers)
                return header_line + "\n" + "\n".join(lines)
        return ""
